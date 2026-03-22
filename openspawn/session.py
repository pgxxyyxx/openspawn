from __future__ import annotations

from dataclasses import replace
from pathlib import Path

from .ai import ClaudeClient, ClaudeError
from .context import build_system_prompt
from .extractor import extract_file
from .launcher import ensure_folder_launcher
from .scanner import describe_changes, detect_changes, prioritize, scan_folder, should_read_file
from .store import AgentStore, load_global_config, save_global_config
from .types import (
    AgentConfig,
    ArtifactRecord,
    ExtractionResult,
    Fact,
    FileChange,
    FileEntry,
    FileIndex,
    HistoryEntry,
    Memory,
    SessionNote,
)
from .utils import utc_now


class AgentSession:
    MAX_CONVERSATION_MESSAGES = 20
    MAX_CONVERSATION_BYTES = 40_000

    def __init__(self, folder: Path):
        self.folder = folder.resolve()
        self.store = AgentStore(self.folder)
        self.global_config = load_global_config()
        self.config = self.store.load_config()
        self.memory = self.store.load_memory()
        self.index = self.store.load_index()
        self.history = self.store.load_history()
        self.session_notes = self.store.load_session_notes()
        self.artifacts = self.store.load_artifacts()
        self.changes: list[FileChange] = []
        self.conversation: list[dict[str, str]] = []
        self.last_response = ""

    @classmethod
    def spawn(cls, folder: Path) -> "AgentSession":
        session = cls(folder)
        if session.config is None:
            session.config = AgentConfig(name=folder.name, root_path=str(folder.resolve()), created_at=utc_now())
        session.refresh(initial=True)
        session.conversation = []
        ensure_folder_launcher(folder)
        session._save()
        session._update_global_agents()
        return session

    def open_existing(self) -> None:
        if self.config is None:
            raise FileNotFoundError("No OpenSpawn agent exists in this folder yet. Run with --init first.")
        ensure_folder_launcher(self.folder)
        self.refresh(initial=False)
        self.conversation = []
        self._save()
        self._update_global_agents()

    def refresh(self, initial: bool) -> None:
        previous_index = self.index
        new_index = scan_folder(self.folder, self.global_config.limits)
        raw_changes = detect_changes(previous_index, new_index)
        old_entries = {entry.path: entry for entry in previous_index.files}
        merged_entries: list[FileEntry] = []
        for entry in prioritize(new_index.files):
            prior = old_entries.get(entry.path)
            if prior and prior.hash == entry.hash and prior.mod_time == entry.mod_time and prior.size == entry.size and prior.read_status in {"read_full", "read_partial", "read_metadata_only"}:
                merged_entries.append(reuse_entry(entry, prior))
                continue
            if should_read_file(entry, self.global_config.limits, sum(1 for item in merged_entries if item.read_status in {"read_full", "read_partial"})):
                merged_entries.append(self._read_entry(entry))
            else:
                merged_entries.append(entry)
        merged_by_path = {entry.path: entry for entry in merged_entries}
        self.changes = describe_changes(raw_changes, old_entries, merged_by_path)
        self.index = FileIndex(files=sorted(merged_entries, key=lambda item: item.path), scanned_at=utc_now(), partial=new_index.partial)
        self._trim_cache()
        self._write_readme()
        self.add_history("Initial scan" if initial else "Refreshed folder", "scan")

    def read_file_now(self, query: str) -> FileEntry | None:
        entry = self.find_file(query)
        if entry is None or entry.is_dir:
            return None
        updated = self._read_entry(entry)
        self._replace_entry(updated)
        self._trim_cache()
        self._write_readme()
        self._save()
        return updated

    def find_file(self, query: str) -> FileEntry | None:
        lowered = query.lower().strip()
        for entry in self.index.files:
            if entry.path.lower() == lowered or entry.name.lower() == lowered:
                return entry
        for entry in self.index.files:
            if lowered in entry.path.lower():
                return entry
        return None

    def add_fact(self, text: str, source: str = "user") -> None:
        self.memory.facts.append(Fact(text=text, created_at=utc_now(), source=source))
        self.add_history(f"Remembered: {text}", "memory")
        self._write_readme()
        self.save()

    def chat(self, client: ClaudeClient, question: str, skill_prompt: str | None = None) -> str:
        prompt, stats = self._prepare_chat_context(question, skill_prompt)
        self.conversation.append({"role": "user", "content": question})
        self._trim_conversation()
        try:
            answer = client.chat_multi(prompt, self.conversation)
        except Exception:
            if self.conversation and self.conversation[-1]["role"] == "user":
                self.conversation.pop()
            raise
        if stats["included_files"] < stats["total_files"]:
            answer = (
                f"{answer}\n\nContext note: Context includes {stats['included_files']} "
                f"of {stats['total_files']} files."
            )
        self.conversation.append({"role": "assistant", "content": answer})
        self._trim_conversation()
        self.last_response = answer
        self.add_history(f"Asked: {question[:80]}", "chat")
        self.save()
        return answer

    def chat_stream(self, client: ClaudeClient, question: str, skill_prompt: str | None = None):
        prompt, stats = self._prepare_chat_context(question, skill_prompt)
        self.conversation.append({"role": "user", "content": question})
        self._trim_conversation()
        parts: list[str] = []
        try:
            for chunk in client.chat_multi_stream(prompt, self.conversation):
                parts.append(chunk)
                yield chunk
        except Exception:
            if self.conversation and self.conversation[-1]["role"] == "user":
                self.conversation.pop()
            raise
        answer = "".join(parts).strip()
        if stats["included_files"] < stats["total_files"]:
            note = (
                f"\n\nContext note: Context includes {stats['included_files']} "
                f"of {stats['total_files']} files."
            )
            answer = f"{answer}{note}"
            yield note
        # Store final answer (with context note) in conversation buffer
        self.conversation.append({"role": "assistant", "content": answer})
        self._trim_conversation()
        self.last_response = answer
        self.add_history(f"Asked: {question[:80]}", "chat")
        self.save()

    def _prepare_chat_context(self, question: str, skill_prompt: str | None = None):
        promoted = self._promoted_paths(question)
        prompt, stats = build_system_prompt(
            self.config,
            self.index.files,
            self.memory,
            self.changes,
            self.history,
            self.session_notes,
            self.artifacts,
            promoted,
            skill_prompt,
        )
        return prompt, stats

    def _promoted_paths(self, question: str) -> set[str]:
        """Determine which files get Level 2 (content) inclusion.

        Uses deterministic cheap signals:
          1. Changed files
          2. Files explicitly mentioned in the question
          3. Files referenced in recent chat history
        """
        paths: set[str] = set()
        # Changed files
        for change in self.changes:
            paths.add(change.path)
        # Files mentioned in the question
        self._match_file_references(question, paths)
        # Files mentioned in recent conversation (last 6 messages)
        for msg in self.conversation[-6:]:
            self._match_file_references(msg.get("content", ""), paths)
            if len(paths) >= 12:
                break
        return paths

    def _match_file_references(self, text: str, paths: set[str]) -> None:
        """Find file references in text. Requires the name include its extension
        (e.g. 'report.pdf') to avoid matching common words like 'status' or 'data'."""
        text_lower = text.lower()
        for entry in self.index.files:
            # Full path match — always safe
            if entry.path.lower() in text_lower:
                paths.add(entry.path)
                continue
            # Name with extension match — 'report.pdf' in text, not bare 'report'
            name_lower = entry.name.lower()
            if "." in name_lower and name_lower in text_lower:
                paths.add(entry.path)

    def status_lines(self) -> list[str]:
        indexed = len(self.index.files)
        readable = sum(1 for entry in self.index.files if entry.read_status in {"read_full", "read_partial"})
        metadata_only = sum(1 for entry in self.index.files if entry.read_status == "read_metadata_only")
        errors = sum(1 for entry in self.index.files if entry.error)
        lines = [
            f"{self.config.name} — {indexed} files indexed",
            f"{readable} files read, {metadata_only} metadata-only, {errors} with errors",
        ]
        if self.index.partial:
            lines.append("Index is partial because folder limits were reached.")
        return lines

    def capture_session_note(self, client: ClaudeClient | None = None) -> SessionNote:
        remembered_facts = [fact.text for fact in self.memory.facts[-5:]]
        open_threads: list[str] = []
        decisions: list[str] = []
        if client is not None:
            open_threads, decisions = self._generate_session_handoff(client)
        if not open_threads:
            open_threads = self._fallback_open_threads()
        note = SessionNote(
            created_at=utc_now(),
            remembered_facts=remembered_facts,
            open_threads=open_threads[:5],
            decisions=decisions[:5],
        )
        self.session_notes.append(note)
        self.add_history("Captured session note", "session")
        self._write_done_note(note)
        self.save()
        return note

    def save_response(self, filename: str) -> Path:
        content = self.last_response
        if not content:
            raise ValueError("No response to save yet.")
        path = self.folder / filename
        if path.suffix == ".docx":
            _save_as_docx(content, path)
        else:
            path.write_text(content, encoding="utf-8")
        self.artifacts.append(
            ArtifactRecord(
                filename=path.name,
                artifact_type=_artifact_type_for_path(path),
                created_at=utc_now(),
            )
        )
        self.add_history(f"Saved response to {filename}", "save")
        self.save()
        return path

    def add_history(self, action: str, entry_type: str, details: str = "") -> None:
        self.history.append(HistoryEntry(time=utc_now(), action=action, entry_type=entry_type, details=details))

    def save(self) -> None:
        save_global_config(self.global_config)
        self._save()

    def close(self) -> None:
        self.add_history("Session closed", "system")
        self.save()

    def _read_entry(self, entry: FileEntry) -> FileEntry:
        result = extract_file(self.folder / entry.path, entry, self.global_config.limits)
        return replace(
            entry,
            read_status=result.status,
            parser=result.parser,
            bytes_read=result.bytes_read,
            content_text=result.content_text[: self.global_config.limits.max_text_bytes],
            structured_summary=result.structured_summary,
            warnings=result.warnings,
            error=result.error,
            summary=summarize_extraction(entry, result),
        )

    def _replace_entry(self, updated: FileEntry) -> None:
        self.index = replace(self.index, files=[updated if entry.path == updated.path else entry for entry in self.index.files])

    def _trim_cache(self) -> None:
        max_per_file = self.global_config.limits.max_text_bytes
        for entry in self.index.files:
            if len(entry.content_text.encode("utf-8")) > max_per_file:
                entry.content_text = trim_to_bytes(entry.content_text, max_per_file)
        total = sum(len(entry.content_text.encode("utf-8")) for entry in self.index.files)
        cap = self.global_config.limits.max_folder_cache_bytes
        if total <= cap:
            return
        for entry in sorted(self.index.files, key=lambda item: (item.priority, len(item.content_text))):
            if total <= cap:
                break
            if not entry.content_text:
                continue
            total -= len(entry.content_text.encode("utf-8"))
            entry.content_text = ""
            entry.read_status = "read_metadata_only"
            entry.warnings = list(entry.warnings) + ["Content trimmed for cache limits. Metadata only until re-read."]

    def _save(self) -> None:
        self.store.save_config(self.config)
        self.store.save_memory(self.memory)
        self.store.save_index(self.index)
        self.store.save_history(self.history)
        self.store.save_session_notes(self.session_notes)
        self.store.save_artifacts(self.artifacts)

    def _trim_conversation(self) -> None:
        while self.conversation and (
            len(self.conversation) > self.MAX_CONVERSATION_MESSAGES
            or self._conversation_size_bytes() > self.MAX_CONVERSATION_BYTES
        ):
            delete_count = 2 if len(self.conversation) >= 2 else 1
            del self.conversation[:delete_count]

    def _conversation_size_bytes(self) -> int:
        total = 0
        for message in self.conversation:
            total += len(message.get("content", "").encode("utf-8"))
        return total

    def _write_readme(self) -> None:
        lines = [f"# {self.config.name}", "", f"Last scan: {self.index.scanned_at}", "", "## Status"]
        lines.extend(f"- {line}" for line in self.status_lines())
        lines.extend(["", "## Files"])
        for entry in self.index.files[:100]:
            lines.append(f"- `{entry.path}` — {entry.read_status}{f' — {entry.summary}' if entry.summary else ''}")
        unreadable = [entry for entry in self.index.files if entry.error]
        if unreadable:
            lines.extend(["", "## Warnings"])
            for entry in unreadable[:20]:
                lines.append(f"- `{entry.path}` — {entry.error}")
        if self.memory.facts:
            lines.extend(["", "## Memory"])
            for fact in self.memory.facts[-20:]:
                lines.append(f"- {fact.text}")
        if self.session_notes:
            latest = self.session_notes[-1]
            if latest.open_threads:
                lines.extend(["", "## Open Threads"])
                for item in latest.open_threads[:5]:
                    lines.append(f"- {item}")
        if self.artifacts:
            lines.extend(["", "## Saved Artifacts"])
            for artifact in self.artifacts[-10:]:
                lines.append(f"- `{artifact.filename}` — {artifact.artifact_type} — {artifact.created_at}")
        self.store.write_readme("\n".join(lines) + "\n")

    def _write_done_note(self, note: SessionNote) -> None:
        lines = ["# Session Note", "", f"Created: {note.created_at}", "", "## Remembered Facts"]
        lines.extend(f"- {item}" for item in (note.remembered_facts or ["None yet."]))
        lines.extend(["", "## Open Threads"])
        lines.extend(f"- {item}" for item in (note.open_threads or ["None."]))
        lines.extend(["", "## Decisions"])
        lines.extend(f"- {item}" for item in (note.decisions or ["None."]))
        self.store.write_done_note("\n".join(lines) + "\n")

    def _update_global_agents(self) -> None:
        agents = [agent for agent in self.global_config.agents if agent.get("path") != str(self.folder)]
        agents.append({"path": str(self.folder), "name": self.config.name, "last_opened": utc_now()})
        self.global_config.agents = agents[-100:]
        save_global_config(self.global_config)

    def _generate_session_handoff(self, client: ClaudeClient) -> tuple[list[str], list[str]]:
        promoted = self._promoted_paths("")
        prompt, _stats = build_system_prompt(
            self.config,
            self.index.files,
            self.memory,
            self.changes,
            self.history,
            self.session_notes,
            self.artifacts,
            promoted,
        )
        recent_response = self.last_response.strip()
        if recent_response:
            recent_response = recent_response[:4000]
        user_message_lines = [
            "Create a concise session handoff from the current folder state.",
            "Return exactly these sections and nothing else:",
            "Open Threads:",
            "- ...",
            "Decisions:",
            "- ...",
            "Rules:",
            "- Keep each bullet concrete and grounded in the folder or recent work.",
            "- Use 1-3 bullets per section.",
            "- If there is nothing to list, write '- None'.",
        ]
        if recent_response:
            user_message_lines.extend(["", "Most recent assistant response:", recent_response])
        try:
            response = client.chat(prompt, "\n".join(user_message_lines), max_tokens=600)
        except ClaudeError:
            return [], []
        sections = _parse_handoff_sections(response)
        return sections["open_threads"], sections["decisions"]

    def _fallback_open_threads(self) -> list[str]:
        items: list[str] = []
        for entry in reversed(self.history):
            if entry.entry_type == "chat" and entry.action.startswith("Asked: "):
                question = entry.action.removeprefix("Asked: ").strip()
                if question and question not in items:
                    items.append(question)
            if len(items) >= 3:
                break
        return items


def reuse_entry(current: FileEntry, previous: FileEntry) -> FileEntry:
    return replace(
        current,
        read_status=previous.read_status,
        bytes_read=previous.bytes_read,
        parser=previous.parser,
        summary=previous.summary,
        warnings=list(previous.warnings),
        error=previous.error,
        content_text=previous.content_text,
        structured_summary=dict(previous.structured_summary),
    )


def summarize_extraction(entry: FileEntry, result: ExtractionResult) -> str:
    if result.error:
        return result.error
    kind = result.structured_summary.get("kind")
    if kind == "table":
        return f"Structured table file with {len(result.structured_summary.get('columns', []))} columns."
    if kind == "workbook":
        return f"Workbook with {len(result.structured_summary.get('sheets', []))} sampled sheets."
    if kind == "pdf":
        return "PDF text extracted." if result.content_text else "PDF metadata only."
    if kind == "notebook":
        return f"Notebook with {result.structured_summary.get('cell_count', 0)} cells."
    if kind == "document":
        return f"Word document with {result.structured_summary.get('paragraph_count', 0)} paragraphs."
    if kind == "presentation":
        return f"Presentation with {result.structured_summary.get('slide_count', 0)} slides."
    return f"Read {entry.name}." if result.content_text else ""


def trim_to_bytes(text: str, max_bytes: int) -> str:
    return text.encode("utf-8")[:max_bytes].decode("utf-8", errors="ignore")


def _artifact_type_for_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return "docx"
    if suffix == ".md":
        return "markdown"
    if suffix == ".txt":
        return "text"
    return suffix.lstrip(".") or "file"


def _parse_handoff_sections(text: str) -> dict[str, list[str]]:
    mapping = {
        "open threads:": "open_threads",
        "decisions:": "decisions",
    }
    sections = {value: [] for value in mapping.values()}
    current: str | None = None
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        key = mapping.get(line.lower())
        if key is not None:
            current = key
            continue
        if current is None or not line.startswith("-"):
            continue
        item = line[1:].strip()
        if not item or item.lower() == "none":
            continue
        sections[current].append(item)
    return sections


def _save_as_docx(content: str, path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        path.with_suffix(".md").write_text(content, encoding="utf-8")
        return
    doc = Document()
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)
    doc.save(str(path))
